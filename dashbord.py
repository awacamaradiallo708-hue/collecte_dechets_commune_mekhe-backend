"""
DASHBOARD DE SUIVI DES COLLECTES - COMMUNE DE MÉKHÉ
Version adaptée à la structure de la base de données
- Suivi quotidien, hebdomadaire, mensuel, annuel
- Graphiques interactifs
- Export Excel et Word
- Panneau d'administration
- Rapports imprimables en PDF (HTML)
- Unités : m³
- Carte interactive avec couleur par équipe et taille par collecte
"""

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import date, datetime, timedelta
from sqlalchemy import create_engine, text
import os
from docx import Document
import calendar
from math import radians, sin, cos, sqrt, atan2

st.set_page_config(
    page_title="Dashboard Collecte - Mékhé",
    page_icon="📊",
    layout="wide"
)

# ==================== STYLE CSS ====================
st.markdown("""
    <style>
    .main-header {
        background: linear-gradient(135deg, #2E7D32 0%, #1B5E20 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 1rem;
    }
    .info-box {
        background: #e3f2fd;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #2196F3;
        margin: 1rem 0;
    }
    .admin-box {
        background: #fff3e0;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #FF9800;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-header"><h1>📊 Dashboard de Suivi des Collectes</h1><p>Commune de Mékhé | Suivi temps réel | Unité : m³</p></div>', unsafe_allow_html=True)

# ==================== CONNEXION BASE DE DONNÉES ====================
DATABASE_URL = os.getenv("DATABASE_URL")
if not DATABASE_URL:
    st.error("❌ Configuration base de données manquante")
    st.stop()

engine = create_engine(DATABASE_URL, pool_pre_ping=True)

# ==================== FONCTIONS UTILITAIRES ====================
def haversine(lat1, lon1, lat2, lon2):
    """Distance en km entre deux points GPS"""
    R = 6371
    lat1, lon1, lat2, lon2 = map(radians, [lat1, lon1, lat2, lon2])
    dlat = lat2 - lat1
    dlon = lon2 - lon1
    a = sin(dlat/2)**2 + cos(lat1) * cos(lat2) * sin(dlon/2)**2
    c = 2 * atan2(sqrt(a), sqrt(1-a))
    return R * c

def load_all_data():
    """Charge les données depuis la base avec les noms de colonnes originaux"""
    with engine.connect() as conn:
        # Tournées
        query_tournees = text("""
            SELECT 
                t.id,
                t.date_tournee,
                t.agent_nom,
                q.nom as quartier,
                e.nom as equipe,
                t.volume_collecte1,
                t.volume_collecte2,
                t.volume_m3,
                t.distance_parcourue_km,
                t.heure_depot_depart,
                t.heure_retour_depot,
                t.created_at,
                (SELECT COUNT(*) FROM points_arret WHERE tournee_id = t.id) as nb_points
            FROM tournees t
            JOIN quartiers q ON t.quartier_id = q.id
            JOIN equipes e ON t.equipe_id = e.id
            WHERE t.statut = 'termine'
            ORDER BY t.date_tournee DESC
        """)
        df = pd.read_sql(query_tournees, conn)
        if not df.empty:
            df['date'] = pd.to_datetime(df['date_tournee'])
            df['semaine'] = df['date'].dt.isocalendar().week
            df['annee'] = df['date'].dt.year
            df['mois'] = df['date'].dt.month
            df.rename(columns={'agent_nom': 'agent', 'distance_parcourue_km': 'distance'}, inplace=True)
        
        # Points GPS
        query_points = text("""
            SELECT 
                pa.tournee_id,
                pa.heure,
                pa.type_point,
                pa.lat,
                pa.lon,
                pa.collecte_numero,
                pa.description,
                q.nom as quartier,
                e.nom as equipe,
                t.date_tournee
            FROM points_arret pa
            JOIN tournees t ON pa.tournee_id = t.id
            JOIN quartiers q ON t.quartier_id = q.id
            JOIN equipes e ON t.equipe_id = e.id
            WHERE pa.lat IS NOT NULL
            ORDER BY t.date_tournee DESC, pa.heure
        """)
        df_points = pd.read_sql(query_points, conn)
        return df, df_points

def formater_duree(minutes):
    if minutes <= 0:
        return "0 min"
    h = int(minutes // 60)
    m = int(minutes % 60)
    return f"{h}h {m}min" if h > 0 else f"{m}min"

def generer_rapport_html(df, periode_nom):
    """Génère un rapport HTML complet pour la période sélectionnée"""
    if df.empty:
        return "<p>Aucune donnée pour cette période.</p>"
    
    total_volume = df['volume_m3'].sum()
    total_tonnes = total_volume * 0.8
    total_distance = df['distance'].sum()
    nb_tournees = len(df)
    nb_quartiers = df['quartier'].nunique()
    nb_agents = df['agent'].nunique()
    top_quartier = df.groupby('quartier')['volume_m3'].sum().idxmax() if not df.empty else "N/A"
    top_agent = df.groupby('agent')['volume_m3'].sum().idxmax() if not df.empty else "N/A"
    
    evol_jour = df.groupby('date')['volume_m3'].sum().reset_index()
    fig1 = px.line(evol_jour, x='date', y='volume_m3', title="Volume collecté par jour (m³)", markers=True)
    fig1.update_layout(height=400)
    graph1_html = fig1.to_html(include_plotlyjs='cdn', div_id="graph1")
    
    top_quartiers = df.groupby('quartier')['volume_m3'].sum().sort_values(ascending=False)
    fig2 = px.bar(x=top_quartiers.values, y=top_quartiers.index, orientation='h',
                  title="Volume total par quartier (m³)", text=top_quartiers.values)
    fig2.update_traces(texttemplate='%{text:.1f} m³', textposition='outside')
    fig2.update_layout(height=400)
    graph2_html = fig2.to_html(include_plotlyjs='cdn', div_id="graph2")
    
    fig3 = px.pie(values=top_quartiers.values, names=top_quartiers.index, title="Répartition des volumes")
    fig3.update_traces(textinfo='percent+label')
    graph3_html = fig3.to_html(include_plotlyjs='cdn', div_id="graph3")
    
    tableau_quartiers = df.groupby('quartier').agg({
        'volume_m3': 'sum',
        'distance': 'sum',
        'id': 'count'
    }).round(2)
    tableau_quartiers.columns = ['Volume (m³)', 'Distance (km)', 'Collectes']
    tableau_quartiers['Tonnes'] = (tableau_quartiers['Volume (m³)'] * 0.8).round(1)
    tableau_quartiers = tableau_quartiers.sort_values('Volume (m³)', ascending=False)
    
    tableau_agents = df.groupby('agent').agg({
        'volume_m3': 'sum',
        'id': 'count'
    }).round(2)
    tableau_agents.columns = ['Volume (m³)', 'Collectes']
    tableau_agents['Tonnes'] = (tableau_agents['Volume (m³)'] * 0.8).round(1)
    tableau_agents = tableau_agents.sort_values('Volume (m³)', ascending=False)
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Rapport collectes - {periode_nom}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
            h1 {{ color: #2E7D32; border-bottom: 2px solid #2E7D32; padding-bottom: 10px; }}
            h2 {{ color: #1B5E20; margin-top: 30px; }}
            .header {{ text-align: center; margin-bottom: 30px; }}
            .summary {{
                background: #f8f9fa;
                padding: 15px;
                border-radius: 8px;
                margin: 20px 0;
                display: flex;
                flex-wrap: wrap;
                justify-content: space-between;
            }}
            .metric {{
                flex: 1;
                text-align: center;
                padding: 10px;
                border-right: 1px solid #ddd;
            }}
            .metric:last-child {{ border-right: none; }}
            .metric-value {{ font-size: 24px; font-weight: bold; color: #2E7D32; }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{ background-color: #2E7D32; color: white; }}
            tr:nth-child(even) {{ background-color: #f2f2f2; }}
            .footer {{
                margin-top: 40px;
                text-align: center;
                font-size: 12px;
                color: #666;
                border-top: 1px solid #ddd;
                padding-top: 20px;
            }}
            @media print {{
                body {{ margin: 0; }}
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>📊 Rapport de collecte des déchets</h1>
            <p>Commune de Mékhé – Période : {periode_nom}</p>
            <p>Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}</p>
        </div>
        
        <div class="summary">
            <div class="metric">
                <div class="metric-value">{total_volume:.1f} m³</div>
                <div>Volume total collecté</div>
                <small>({total_tonnes:.1f} tonnes)</small>
            </div>
            <div class="metric">
                <div class="metric-value">{total_distance:.1f} km</div>
                <div>Distance parcourue</div>
            </div>
            <div class="metric">
                <div class="metric-value">{nb_tournees}</div>
                <div>Tournées effectuées</div>
            </div>
            <div class="metric">
                <div class="metric-value">{nb_quartiers}</div>
                <div>Quartiers visités</div>
            </div>
            <div class="metric">
                <div class="metric-value">{nb_agents}</div>
                <div>Agents actifs</div>
            </div>
        </div>
        
        <div class="info">
            <p><strong>🏆 Quartier le plus productif :</strong> {top_quartier}</p>
            <p><strong>👤 Agent le plus performant :</strong> {top_agent}</p>
            <p><strong>⚡ Efficacité globale :</strong> {(total_distance / total_volume) if total_volume>0 else 0:.2f} km/m³</p>
        </div>
        
        <h2>📈 Évolution quotidienne</h2>
        {graph1_html}
        
        <h2>🏘️ Répartition par quartier</h2>
        {graph2_html}
        {graph3_html}
        
        <h2>📊 Tableau des quartiers</h2>
        {tableau_quartiers.to_html()}
        
        <h2>👥 Performance des agents</h2>
        {tableau_agents.to_html()}
        
        <div class="footer">
            <p>Rapport généré automatiquement par le système de suivi des collectes – Commune de Mékhé</p>
            <p>Les volumes sont exprimés en m³ (1 m³ ≈ 0,8 tonne).</p>
        </div>
    </body>
    </html>
    """
    return html

def generer_rapport_docx(df, periode_nom):
    # This function is now defined in application_agent_final.py and will be removed from here.
    """Génère un rapport Word (.docx) complet pour la période sélectionnée"""
    document = Document()
    document.add_heading(f'Rapport de Collecte des Déchets - Mékhé', 0)
    document.add_paragraph(f'Période : {periode_nom}')
    document.add_paragraph(f'Généré le : {datetime.now().strftime("%d/%m/%Y à %H:%M")}')

    # Section Synthèse
    document.add_heading('1. Synthèse de la période', level=1)
    total_vol = df['volume_m3'].sum()
    total_t = total_vol * 0.8
    
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Indicateur'
    hdr_cells[1].text = 'Valeur'
    
    metrics = [
        ('Volume total collecté', f"{total_vol:.2f} m³"),
        ('Poids total estimé', f"{total_t:.2f} Tonnes"),
        ('Nombre de tournées', str(len(df))),
        ('Distance parcourue', f"{df['distance'].sum():.2f} km"),
        ('Nombre de quartiers couverts', str(df['quartier'].nunique()))
    ]
    
    for item, val in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = val

    # Section Production par Quartier
    document.add_heading('2. Production par Quartier', level=1)
    q_stats = df.groupby('quartier').agg({
        'volume_m3': 'sum',
        'id': 'count'
    }).sort_values('volume_m3', ascending=False)
    
    table_q = document.add_table(rows=1, cols=3)
    hdr_q = table_q.rows[0].cells
    hdr_q[0].text = 'Quartier'
    hdr_q[1].text = 'Volume (m³)'
    hdr_q[2].text = 'Nbre Tournées'
    
    for quartier, row in q_stats.iterrows():
        row_q = table_q.add_row().cells
        row_q[0].text = str(quartier)
        row_q[1].text = f"{row['volume_m3']:.2f}"
        row_q[2].text = str(int(row['id']))

    # Section Performance Agents
    document.add_heading('3. Performance des Agents', level=1)
    a_stats = df.groupby('agent')['volume_m3'].sum().sort_values(ascending=False)
    for agent, vol in a_stats.items():
        document.add_paragraph(f'- {agent} : {vol:.2f} m³ collectés', style='List Bullet')

    document.add_paragraph('\nNote: Les tonnes sont calculées sur une base de densité moyenne de 0.8.')
    
    buffer = BytesIO()
    # Ensure Document is not None before calling save
    document.save(buffer)
    return buffer.getvalue()

# ==================== PAGE ADMINISTRATION ====================
def show_admin_panel():
    """Panneau d'administration pour l'équipe technique"""
    st.markdown('<div class="admin-box">🔧 <strong>Panneau d\'administration</strong> - Accès réservé à l\'équipe technique</div>', unsafe_allow_html=True)
    
    tab1, tab2, tab3, tab4 = st.tabs([
        "📊 Statistiques globales", 
        "👥 Gestion des agents", 
        "🏘️ Gestion des quartiers",
        "📁 Exports et sauvegarde"
    ])
    
    with tab1:
        st.subheader("📊 Statistiques globales")
        with engine.connect() as conn:
            total = conn.execute(text("SELECT COUNT(*) FROM tournees WHERE statut = 'termine'")).scalar()
            st.metric("📦 Total collectes", total)
            volume = conn.execute(text("SELECT SUM(volume_m3) FROM tournees WHERE statut = 'termine'")).scalar()
            st.metric("📊 Volume total", f"{volume:.1f} m³" if volume else "0")
            agents = conn.execute(text("SELECT COUNT(DISTINCT agent_nom) FROM tournees WHERE statut = 'termine'")).scalar()
            st.metric("👥 Agents actifs", agents)
            derniere = conn.execute(text("SELECT MAX(date_tournee) FROM tournees WHERE statut = 'termine'")).scalar()
            derniere_str = derniere.strftime("%d/%m/%Y") if derniere else "Aucune"
            st.metric("📅 Dernière collecte", derniere_str)
        
        col1, col2 = st.columns(2)
        with col1:
            with engine.connect() as conn:
                df_activite = pd.read_sql("""
                    SELECT date_tournee, COUNT(*) as nb_collectes, SUM(volume_m3) as volume
                    FROM tournees WHERE statut = 'termine'
                    GROUP BY date_tournee ORDER BY date_tournee DESC LIMIT 30
                """, conn)
                if not df_activite.empty:
                    fig = px.bar(df_activite, x='date_tournee', y='nb_collectes', title="Activité des 30 derniers jours")
                    st.plotly_chart(fig, use_container_width=True)
        with col2:
            with engine.connect() as conn:
                df_volume = pd.read_sql("""
                    SELECT date_tournee, SUM(volume_m3) as volume_total
                    FROM tournees WHERE statut = 'termine'
                    GROUP BY date_tournee ORDER BY date_tournee DESC LIMIT 30
                """, conn)
                if not df_volume.empty:
                    fig2 = px.line(df_volume, x='date_tournee', y='volume_total', title="Volume collecté des 30 derniers jours (m³)", markers=True)
                    st.plotly_chart(fig2, use_container_width=True)
        
        with engine.connect() as conn:
            df_mois = pd.read_sql("""
                SELECT EXTRACT(YEAR FROM date_tournee) as annee, EXTRACT(MONTH FROM date_tournee) as mois,
                       COUNT(*) as nb_collectes, SUM(volume_m3) as volume_total
                FROM tournees WHERE statut = 'termine'
                GROUP BY annee, mois ORDER BY annee DESC, mois DESC LIMIT 12
            """, conn)
            if not df_mois.empty:
                df_mois['periode'] = df_mois['annee'].astype(int).astype(str) + '-' + df_mois['mois'].astype(int).astype(str).str.zfill(2)
                fig3 = px.bar(df_mois, x='periode', y='volume_total', title="Volume collecté par mois (m³)")
                st.plotly_chart(fig3, use_container_width=True)
    
    with tab2:
        st.subheader("👥 Liste des agents")
        with engine.connect() as conn:
            agents_df = pd.read_sql("""
                SELECT agent_nom, COUNT(*) as nb_collectes, COALESCE(SUM(volume_m3),0) as volume_total_m3,
                       COALESCE(AVG(volume_m3),0) as volume_moyen_m3,
                       COALESCE(SUM(distance_parcourue_km),0) as distance_totale_km,
                       MAX(date_tournee) as derniere_activite
                FROM tournees WHERE statut = 'termine'
                GROUP BY agent_nom ORDER BY volume_total_m3 DESC
            """, conn)
            if not agents_df.empty:
                agents_df = agents_df.fillna(0)
                agents_df['derniere_activite'] = pd.to_datetime(agents_df['derniere_activite']).dt.strftime('%d/%m/%Y')
                agents_df.columns = ['Agent', 'Nb collectes', 'Volume total (m³)', 'Volume moyen (m³)', 'Distance totale (km)', 'Dernière activité']
                st.dataframe(agents_df, use_container_width=True)
                col1, col2 = st.columns(2)
                with col1:
                    top_agents = agents_df.head(10)
                    fig_agents = px.bar(top_agents, x='Agent', y='Volume total (m³)', title="Top 10 agents (volume collecté)")
                    st.plotly_chart(fig_agents, use_container_width=True)
                with col2:
                    csv = agents_df.to_csv(index=False).encode('utf-8')
                    st.download_button("📥 Exporter la liste des agents (CSV)", csv, f"agents_{datetime.now().strftime('%Y%m%d')}.csv")
            else:
                st.info("Aucun agent enregistré")
    
    with tab3:
        st.subheader("🏘️ Performance par quartier")
        with engine.connect() as conn:
            quartiers_df = pd.read_sql("""
                SELECT q.nom as quartier, COALESCE(q.population,0) as population,
                       COUNT(t.id) as nb_collectes, COALESCE(SUM(t.volume_m3),0) as volume_total_m3,
                       COALESCE(AVG(t.volume_m3),0) as volume_moyen_m3,
                       COALESCE(SUM(t.distance_parcourue_km),0) as distance_totale_km
                FROM quartiers q
                LEFT JOIN tournees t ON q.id = t.quartier_id AND t.statut = 'termine'
                GROUP BY q.nom, q.population ORDER BY volume_total_m3 DESC
            """, conn)
            if not quartiers_df.empty:
                quartiers_df = quartiers_df.fillna(0)
                quartiers_df['m3_par_habitant'] = quartiers_df.apply(
                    lambda row: row['volume_total_m3'] / row['population'] if row['population'] > 0 else 0, axis=1)
                quartiers_df_display = quartiers_df.copy()
                quartiers_df_display.columns = ['Quartier', 'Population', 'Nb collectes', 'Volume total (m³)', 'Volume moyen (m³)', 'Distance totale (km)', 'm³ par habitant']
                st.dataframe(quartiers_df_display, use_container_width=True)
                
                col1, col2 = st.columns(2)
                with col1:
                    quartiers_filtres = quartiers_df[quartiers_df['volume_total_m3'] > 0].head(10)
                    if not quartiers_filtres.empty:
                        fig_volume = px.bar(quartiers_filtres, x='quartier', y='volume_total_m3', title="Top quartiers - Volume total (m³)")
                        st.plotly_chart(fig_volume, use_container_width=True)
                with col2:
                    quartiers_filtres_hab = quartiers_df[quartiers_df['m3_par_habitant'] > 0].head(10)
                    if not quartiers_filtres_hab.empty:
                        fig_habitant = px.bar(quartiers_filtres_hab, x='quartier', y='m3_par_habitant', title="Volume par habitant (m³/hab)")
                        st.plotly_chart(fig_habitant, use_container_width=True)
                
                st.subheader("📊 Relation Volume collecté vs Population")
                quartiers_scatter = quartiers_df[(quartiers_df['population'] > 0) & (quartiers_df['volume_total_m3'] > 0)]
                if not quartiers_scatter.empty:
                    fig_scatter = px.scatter(quartiers_scatter, x='population', y='volume_total_m3', size='volume_total_m3', text='quartier')
                    fig_scatter.update_traces(textposition='top center')
                    st.plotly_chart(fig_scatter, use_container_width=True)
                
                csv = quartiers_df_display.to_csv(index=False).encode('utf-8')
                st.download_button("📥 Exporter les données des quartiers (CSV)", csv, f"quartiers_{datetime.now().strftime('%Y%m%d')}.csv")
            else:
                st.info("Aucune donnée disponible pour les quartiers")
    
    with tab4:
        st.subheader("📁 Exports et sauvegarde")
        st.markdown("""
        <div class="info-box">
        <strong>📋 Instructions :</strong><br>
        - Export complet : Toutes les données de la base<br>
        - Export période : Données selon la période sélectionnée<br>
        - Sauvegarde : Export hebdomadaire recommandé
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("#### 📊 Export complet")
            if st.button("📥 Exporter TOUTES les données", use_container_width=True):
                with engine.connect() as conn:
                    df_all = pd.read_sql("""
                        SELECT t.date_tournee, t.agent_nom, q.nom as quartier, e.nom as equipe,
                               t.volume_collecte1, t.volume_collecte2, t.volume_m3,
                               t.distance_parcourue_km, t.heure_depot_depart, t.heure_retour_depot,
                               t.created_at, (SELECT COUNT(*) FROM points_arret WHERE tournee_id = t.id) as nb_points
                        FROM tournees t
                        JOIN quartiers q ON t.quartier_id = q.id
                        JOIN equipes e ON t.equipe_id = e.id
                        WHERE t.statut = 'termine'
                        ORDER BY t.date_tournee DESC
                    """, conn)
                    if not df_all.empty:
                        output = BytesIO()
                        with pd.ExcelWriter(output, engine='openpyxl') as writer:
                            df_all.to_excel(writer, sheet_name="Toutes les collectes", index=False)
                            synth = df_all.groupby('quartier').agg({'volume_m3': 'sum', 'distance_parcourue_km': 'sum'}).round(2)
                            synth.to_excel(writer, sheet_name="Synthèse par quartier")
                            df_all['mois'] = df_all['date_tournee'].dt.strftime('%Y-%m')
                            evol = df_all.groupby('mois').agg({'volume_m3': 'sum'}).round(2)
                            evol.to_excel(writer, sheet_name="Évolution mensuelle")
                        st.download_button("📥 Télécharger l'export complet (Excel)", data=output.getvalue(),
                                           file_name=f"export_complet_{datetime.now().strftime('%Y%m%d')}.xlsx")
                    else:
                        st.warning("Aucune donnée à exporter")
        with col2:
            st.markdown("#### 🗄️ Sauvegarde de la base")
            st.info("""
            **🔧 Commandes utiles pour la sauvegarde :**
            
            Via Neon.tech :
            1. Aller sur https://neon.tech
            2. Ouvrir le projet
            3. Aller dans "Backups"
            4. Créer une sauvegarde manuelle
            """)
        
        st.markdown("---")
        st.markdown("#### 📈 Rapport de synthèse")
        date_debut = st.date_input("Date début", value=date.today() - timedelta(days=30))
        date_fin = st.date_input("Date fin", value=date.today())
        if st.button("📊 Générer rapport de synthèse", use_container_width=True):
            with engine.connect() as conn:
                df_periode = pd.read_sql(f"""
                    SELECT t.date_tournee, t.agent_nom, q.nom as quartier, t.volume_m3, t.distance_parcourue_km
                    FROM tournees t JOIN quartiers q ON t.quartier_id = q.id
                    WHERE t.statut = 'termine' AND t.date_tournee BETWEEN '{date_debut}' AND '{date_fin}'
                    ORDER BY t.date_tournee
                """, conn)
                if not df_periode.empty:
                    stats = {
                        "nb_tournees": len(df_periode),
                        "volume_total": df_periode['volume_m3'].sum(),
                        "distance_total": df_periode['distance_parcourue_km'].sum(),
                        "nb_quartiers": df_periode['quartier'].nunique(),
                        "nb_agents": df_periode['agent_nom'].nunique(),
                        "top_quartier": df_periode.groupby('quartier')['volume_m3'].sum().idxmax()
                    }
                    html_content = generer_rapport_html(df_periode, f"{date_debut} au {date_fin}")
                    st.download_button("📥 Télécharger le rapport Word", data=html_content,
                                       file_name=f"rapport_synthese_{date_debut}_{date_fin}.html")
                else:
                    st.warning("Aucune donnée pour cette période")

# ==================== CHARGEMENT DES DONNÉES ====================
with st.spinner("Chargement des données..."):
    df_tournees, df_points = load_all_data()

if df_tournees.empty:
    st.warning("⚠️ Aucune donnée disponible. Les agents doivent d'abord enregistrer des collectes.")
    st.info("""
    **Pour commencer :**
    1. Les agents doivent utiliser l'application de collecte
    2. Enregistrer leurs tournées
    3. Les données apparaîtront ici automatiquement
    """)
    st.stop()

# ==================== BARRE LATÉRALE ====================
with st.sidebar:
    st.header("🎛️ Filtres")
    periode = st.selectbox("Période d'analyse", ["Aujourd'hui", "Cette semaine", "Ce mois", "Personnalisé"])
    
    if periode == "Aujourd'hui":
        date_filter = st.date_input("Date", value=date.today())
        df_filtered = df_tournees[df_tournees['date'].dt.date == date_filter]
        periode_nom = date_filter.strftime("%d/%m/%Y")
    elif periode == "Cette semaine":
        today = date.today()
        start_of_week = today - timedelta(days=today.weekday())
        end_of_week = start_of_week + timedelta(days=6)
        df_filtered = df_tournees[(df_tournees['date'].dt.date >= start_of_week) & (df_tournees['date'].dt.date <= end_of_week)]
        periode_nom = f"Semaine du {start_of_week.strftime('%d/%m')} au {end_of_week.strftime('%d/%m/%Y')}"
    elif periode == "Ce mois":
        today = date.today()
        start_of_month = today.replace(day=1)
        df_filtered = df_tournees[df_tournees['date'].dt.date >= start_of_month]
        periode_nom = f"Mois de {calendar.month_name[today.month]} {today.year}"
    else:
        col1, col2 = st.columns(2)
        with col1:
            date_debut = st.date_input("Date début", value=date.today() - timedelta(days=30))
        with col2:
            date_fin = st.date_input("Date fin", value=date.today())
        df_filtered = df_tournees[(df_tournees['date'].dt.date >= date_debut) & (df_tournees['date'].dt.date <= date_fin)]
        periode_nom = f"Du {date_debut.strftime('%d/%m/%Y')} au {date_fin.strftime('%d/%m/%Y')}"
    
    st.markdown("---")
    quartiers = st.multiselect("Quartiers", df_filtered['quartier'].unique(), default=df_filtered['quartier'].unique())
    agents = st.multiselect("Agents", df_filtered['agent'].unique(), default=df_filtered['agent'].unique())
    df_filtered = df_filtered[df_filtered['quartier'].isin(quartiers)]
    df_filtered = df_filtered[df_filtered['agent'].isin(agents)]

ids_tournees = df_filtered['id'].tolist()
df_points_filtre = df_points[df_points['tournee_id'].isin(ids_tournees)]

# ==================== ONGLETS ====================
tabs = st.tabs(["📈 Tableau de bord", "🥇 Classements", "🗺️ Carte", "📋 Détails", "📊 Rapports"])

# ==================== TAB 1 : TABLEAU DE BORD ====================
with tabs[0]:
    st.subheader(f"📅 Période : {periode_nom}")
    if not df_filtered.empty:
        total_volume = df_filtered['volume_m3'].sum()
        total_distance = df_filtered['distance'].sum()
        nb_tournees = len(df_filtered)
        nb_quartiers = df_filtered['quartier'].nunique()
        nb_agents = df_filtered['agent'].nunique()
        
        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            st.metric("📦 Volume total", f"{total_volume:.1f} m³")
            st.caption(f"≈ {total_volume * 0.8:.0f} tonnes")
        with col2:
            st.metric("📏 Distance totale", f"{total_distance:.1f} km")
        with col3:
            st.metric("🚛 Tournées", nb_tournees)
        with col4:
            st.metric("🏘️ Quartiers", nb_quartiers)
        with col5:
            st.metric("👥 Agents", nb_agents)
        
        if total_volume > 0:
            st.info(f"📊 **Efficacité globale :** {total_distance/total_volume:.2f} km/m³")
        
        st.markdown("---")
        col_charts1, col_charts2 = st.columns(2)
        
        with col_charts1:
            evol_journaliere = df_filtered.groupby('date')['volume_m3'].sum().reset_index()
            fig_evol = px.line(evol_journaliere, x='date', y='volume_m3', title="Production journalière (m³)", markers=True, color_discrete_sequence=['#2E7D32'])
            st.plotly_chart(fig_evol, use_container_width=True)
            
        with col_charts2:
            prod_quartier = df_filtered.groupby('quartier')['volume_m3'].sum().sort_values(ascending=False).reset_index()
            fig_prod = px.bar(prod_quartier, x='quartier', y='volume_m3', 
                             title="Production totale par quartier (m³)",
                             color='volume_m3', color_continuous_scale='Greens')
            st.plotly_chart(fig_prod, use_container_width=True)

        st.markdown("---")
        
        evol_quartier = df_filtered.groupby(['date', 'quartier'])['volume_m3'].sum().reset_index()
        fig_quartier = px.line(evol_quartier, x='date', y='volume_m3', color='quartier', title="Évolution par quartier (m³)", markers=True)
        fig_quartier.update_layout(height=500)
        st.plotly_chart(fig_quartier, use_container_width=True)
    else:
        st.warning("⚠️ Aucune donnée pour la période sélectionnée")

# ==================== TAB 2 : CLASSEMENTS ====================
with tabs[1]:
    st.subheader("🥇 Classements")
    if not df_filtered.empty:
        col1, col2 = st.columns(2)
        with col1:
            top_quartiers = df_filtered.groupby('quartier')['volume_m3'].sum().sort_values(ascending=True)
            fig_quartiers = px.bar(x=top_quartiers.values, y=top_quartiers.index, orientation='h',
                                   title="Volume total par quartier (m³)", text=top_quartiers.values)
            fig_quartiers.update_traces(texttemplate='%{text:.1f} m³', textposition='outside')
            st.plotly_chart(fig_quartiers, use_container_width=True)
        with col2:
            top_agents = df_filtered.groupby('agent')['volume_m3'].sum().sort_values(ascending=True)
            fig_agents = px.bar(x=top_agents.values, y=top_agents.index, orientation='h',
                                title="Volume total par agent (m³)", text=top_agents.values)
            fig_agents.update_traces(texttemplate='%{text:.1f} m³', textposition='outside')
            st.plotly_chart(fig_agents, use_container_width=True)
        
        col1, col2 = st.columns(2)
        with col1:
            fig_pie = px.pie(df_filtered, values='volume_m3', names='quartier', title="Répartition des volumes", hole=0.3)
            st.plotly_chart(fig_pie, use_container_width=True)
        with col2:
            perf_df = df_filtered.groupby('quartier').agg({
                'volume_m3': 'sum', 'distance': 'sum', 'id': 'count', 'nb_points': 'sum'
            }).round(2)
            perf_df.columns = ['Volume (m³)', 'Distance (km)', 'Collectes', 'Points GPS']
            perf_df = perf_df.sort_values('Volume (m³)', ascending=False)
            st.dataframe(perf_df, use_container_width=True)
    else:
        st.info("Aucune donnée")

# ==================== TAB 3 : CARTE ====================
with tabs[2]:
    st.subheader("🗺️ Carte des points de collecte")
    
    if not df_points_filtre.empty:
        if 'equipe' not in df_points_filtre.columns:
            st.warning("Les données ne contiennent pas d'information sur les équipes. Affichage sans couleur.")
            df_points_filtre['equipe'] = "Inconnue"
        
        collecte_filtre = st.radio("Afficher les points de :", ["Toutes", "Collecte 1", "Collecte 2"], horizontal=True)
        if collecte_filtre == "Collecte 1":
            df_carte = df_points_filtre[df_points_filtre['collecte_numero'] == 1]
        elif collecte_filtre == "Collecte 2":
            df_carte = df_points_filtre[df_points_filtre['collecte_numero'] == 2]
        else:
            df_carte = df_points_filtre
        
        if not df_carte.empty:
            noms_points = {
                "depart_depot": "🏭 Départ dépôt", "debut_collecte": "🗑️ Début collecte",
                "fin_collecte": "🗑️ Fin collecte", "depart_decharge": "🚛 Départ décharge",
                "arrivee_decharge": "🏭 Arrivée décharge", "sortie_decharge": "🏭 Sortie décharge",
                "retour_depot": "🏁 Retour dépôt", "point_libre": "📍 Point libre"
            }
            df_carte["nom_affichage"] = df_carte["type_point"].map(noms_points)
            df_carte['type_collecte'] = df_carte['collecte_numero'].apply(lambda x: f"Collecte {x}")
            
            equipes = df_carte['equipe'].unique()
            palette = px.colors.qualitative.Plotly
            color_map = {e: palette[i % len(palette)] for i, e in enumerate(equipes)}
            df_carte['taille_point'] = df_carte['collecte_numero'].apply(lambda x: 10 if x == 1 else 15)
            
            fig = px.scatter_mapbox(
                df_carte, 
                lat="lat", 
                lon="lon",
                color="equipe",
                size="taille_point",
                hover_name="nom_affichage",
                hover_data={"quartier": True, "collecte_numero": True, "heure": True, "type_point": True, "lat": False, "lon": False},
                color_discrete_map=color_map,
                zoom=12, 
                center={"lat": 15.11, "lon": -16.65},
                title="Points GPS (couleur = équipe, taille = collecte 1 (petit) / collecte 2 (grand))",
                height=550
            )
            
            for tid in df_carte['tournee_id'].unique():
                df_tour = df_carte[df_carte['tournee_id'] == tid].sort_values('heure')
                if len(df_tour) > 1:
                    equipe_tour = df_tour.iloc[0]['equipe']
                    couleur_ligne = color_map.get(equipe_tour, 'blue')
                    fig.add_trace(go.Scattermapbox(
                        lat=df_tour['lat'].tolist(),
                        lon=df_tour['lon'].tolist(),
                        mode='lines',
                        line=dict(width=2, color=couleur_ligne),
                        name=f'Trajet {equipe_tour} (tournée {tid})',
                        showlegend=False
                    ))
            
            fig.update_layout(mapbox_style="open-street-map", margin={"r":0,"t":40,"l":0,"b":0})
            st.plotly_chart(fig, use_container_width=True)
            
            st.subheader("📏 Distances entre points consécutifs")
            distances = []
            for tid in df_carte['tournee_id'].unique():
                df_tour = df_carte[df_carte['tournee_id'] == tid].sort_values('heure')
                if len(df_tour) > 1:
                    for i in range(1, len(df_tour)):
                        p1 = df_tour.iloc[i-1]
                        p2 = df_tour.iloc[i]
                        d = haversine(p1['lat'], p1['lon'], p2['lat'], p2['lon'])
                        distances.append({
                            "Tournée": tid,
                            "Équipe": p1['equipe'],
                            "De": p1['nom_affichage'],
                            "À": p2['nom_affichage'],
                            "Distance (km)": round(d, 2)
                        })
            if distances:
                st.dataframe(pd.DataFrame(distances), use_container_width=True)
                st.info(f"**Distance totale calculée :** {sum(d['Distance (km)'] for d in distances):.2f} km")
            else:
                st.info("Aucune distance calculable (moins de deux points par tournée)")
        else:
            st.info("Aucun point GPS pour la collecte sélectionnée")
    else:
        st.info("Aucun point GPS enregistré pour la période")

# ==================== TAB 4 : DÉTAILS ====================
with tabs[3]:
    st.subheader("📋 Détail des collectes")
    if not df_filtered.empty:
        display_df = df_filtered.copy()
        display_df['date'] = display_df['date'].dt.strftime('%d/%m/%Y')
        display_df['volume_m3'] = display_df['volume_m3'].apply(lambda x: f"{x:.1f} m³")
        display_df['distance'] = display_df['distance'].apply(lambda x: f"{x:.1f} km")
        display_df['volume_collecte1'] = display_df['volume_collecte1'].apply(lambda x: f"{x:.1f} m³")
        display_df['volume_collecte2'] = display_df['volume_collecte2'].apply(lambda x: f"{x:.1f} m³")
        
        st.dataframe(
            display_df[['date', 'quartier', 'agent', 'equipe', 'volume_collecte1', 'volume_collecte2', 'volume_m3', 'distance', 'nb_points']],
            use_container_width=True,
            column_config={
                "date": "Date", "quartier": "Quartier", "agent": "Agent", "equipe": "Équipe",
                "volume_collecte1": "Volume 1", "volume_collecte2": "Volume 2",
                "volume_m3": "Volume total", "distance": "Distance", "nb_points": "Points GPS"
            }
        )
        
        st.markdown("---")
        st.subheader("📥 Export des données")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📊 EXPORTER EN EXCEL", use_container_width=True):
                output = BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    display_df.to_excel(writer, sheet_name="Collectes", index=False)
                st.download_button("📥 Télécharger Excel", data=output.getvalue(),
                                   file_name=f"collectes_{datetime.now().strftime('%Y%m%d')}.xlsx")
    else:
        st.info("Aucune donnée pour cette période")

# ==================== TAB 5 : RAPPORTS ====================
with tabs[4]:
    st.subheader("📊 Génération de rapports (PDF imprimable)")
    
    # Créer une liste des semaines disponibles à partir des dates
    df_tournees['date'] = pd.to_datetime(df_tournees['date_tournee'])
    df_tournees['semaine_num'] = df_tournees['date'].dt.isocalendar().week
    df_tournees['annee_num'] = df_tournees['date'].dt.year
    
    # Obtenir toutes les combinaisons année-semaine uniques
    semaines_disponibles = df_tournees[['annee_num', 'semaine_num']].drop_duplicates().sort_values(['annee_num', 'semaine_num']).values.tolist()
    
    col1, col2 = st.columns(2)
    with col1:
        type_rapport = st.selectbox("Type de rapport", ["Hebdomadaire", "Mensuel", "Annuel"])
    
    if type_rapport == "Hebdomadaire":
        col1, col2 = st.columns(2)
        with col1:
            # Sélection de l'année
            annees = sorted(df_tournees['annee'].unique(), reverse=True)
            if not annees:
                annees = [date.today().year]
            annee = st.selectbox("Année", annees)
        with col2:
            # Générer les semaines de l'année (1 à 52) et les marquer comme disponibles ou non
            toutes_semaines = list(range(1, 53))
            semaines_avec_donnees = df_tournees[df_tournees['annee'] == annee]['semaine'].unique()
            
            # Créer des libellés avec indicateur de données
            options_semaines = []
            for s in toutes_semaines:
                if s in semaines_avec_donnees:
                    options_semaines.append(f"Semaine {s} (avec données)")
                else:
                    options_semaines.append(f"Semaine {s} (vide)")
            
            semaine_idx = st.selectbox("Semaine", range(len(options_semaines)), format_func=lambda i: options_semaines[i])
            semaine = toutes_semaines[semaine_idx]
            periode_nom_rapport = f"Semaine {semaine} - {annee}"
            
            # Filtrer les données
            df_rapport = df_tournees[(df_tournees['annee'] == annee) & (df_tournees['semaine'] == semaine)]
        
    elif type_rapport == "Mensuel":
        col1, col2 = st.columns(2)
        with col1:
            annees = sorted(df_tournees['annee'].unique(), reverse=True)
            if not annees:
                annees = [date.today().year]
            annee = st.selectbox("Année", annees)
        with col2:
            mois_options = list(range(1, 13))
            mois_avec_donnees = df_tournees[df_tournees['annee'] == annee]['mois'].unique()
            
            options_mois = []
            for m in mois_options:
                if m in mois_avec_donnees:
                    options_mois.append(f"{calendar.month_name[m]} (avec données)")
                else:
                    options_mois.append(f"{calendar.month_name[m]} (vide)")
            
            mois_idx = st.selectbox("Mois", range(len(options_mois)), format_func=lambda i: options_mois[i])
            mois = mois_options[mois_idx]
            nom_mois = calendar.month_name[mois]
            periode_nom_rapport = f"{nom_mois} {annee}"
            
            df_rapport = df_tournees[(df_tournees['annee'] == annee) & (df_tournees['mois'] == mois)]
    
    else:  # Annuel
        annees = sorted(df_tournees['annee'].unique(), reverse=True)
        if not annees:
            annees = [date.today().year]
        annee = st.selectbox("Année", annees)
        periode_nom_rapport = f"Année {annee}"
        df_rapport = df_tournees[df_tournees['annee'] == annee]
    
    if not df_rapport.empty:
        st.info(f"**{len(df_rapport)} tournée(s)** trouvée(s) pour la période")
        
        # Afficher un résumé rapide
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("📦 Volume total", f"{df_rapport['volume_m3'].sum():.1f} m³")
        with col2:
            st.metric("📏 Distance totale", f"{df_rapport['distance'].sum():.1f} km")
        with col3:
            st.metric("🚛 Nombre de tournées", len(df_rapport))
        
        # Removed Word export button
        html_content = generer_rapport_html(df_rapport, periode_nom_rapport)
        st.download_button(
            label="📥 Télécharger Rapport HTML (.html)",
            data=html_content,
            file_name=f"Rapport_Collecte_{periode_nom_rapport.replace(' ', '_')}.html",
            mime="text/html",
            use_container_width=True
        )
    else:
        st.warning(f"Aucune donnée pour la période sélectionnée ({periode_nom_rapport})")
        st.info("💡 Vous pouvez quand même générer un rapport vide ou sélectionner une autre période.")
